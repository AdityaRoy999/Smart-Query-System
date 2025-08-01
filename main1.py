import os
import re
import json
import tempfile
import time
import io
import streamlit as st
import numpy as np
import docx
import requests
from PyPDF2 import PdfReader

# --- DOCUMENT PARSING ---
def extract_text(file_path, file_type):
    """Extracts text from PDF or DOCX files."""
    if file_type == "application/pdf":
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"Error reading PDF file: {e}")
        return text
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {e}")
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX.")

# --- TEXT CHUNKING ---
def chunk_text(text, max_tokens=200):
    """Splits text into chunks of a maximum token size, respecting sentence boundaries."""
    text = re.sub(r'\s+', ' ', text).strip()
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk.split()) + len(sentence.split()) <= max_tokens:
            current_chunk += sentence + " "
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence + " "
            
    if current_chunk:
        chunks.append(current_chunk.strip())
        
    return chunks

# --- API CALLS (Functions now take the key/URL as arguments) ---
@st.cache_data(show_spinner=False)
def get_embedding(text_chunk, api_key):
    """Generates embedding for a given text chunk using the Gemini API with retry logic."""
    embed_url = f"https://generativelanguage.googleapis.com/v1beta/models/embedding-001:embedContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    body = {"model": "models/embedding-001", "content": {"parts": [{"text": text_chunk}]}}
    
    max_retries = 3
    delay = 1
    for attempt in range(max_retries):
        try:
            resp = requests.post(embed_url, headers=headers, json=body, timeout=60)
            resp.raise_for_status()
            embedding = resp.json().get("embedding", {}).get("values")
            if not embedding:
                raise ValueError("API response did not contain embedding values.")
            return embedding
        except requests.exceptions.RequestException as e:
            if attempt < max_retries - 1:
                time.sleep(delay)
                delay *= 2
            else:
                raise RuntimeError(f"API request failed for single embedding after {max_retries} attempts: {e}")
        except (KeyError, ValueError) as e:
            raise RuntimeError(f"Failed to parse API response: {e}\nResponse: {resp.text}")
    return None

@st.cache_data(show_spinner=False)
def embed_chunks(_chunks, api_key):
    """
    Embeds a list of text chunks in batches for efficiency and reliability.
    The underscore in _chunks is a convention to indicate that this argument is used for caching.
    """
    batch_embed_url = f"https://generativelanguage.googleapis.com/v1beta/models/embedding-001:batchEmbedContents?key={api_key}"
    headers = {"Content-Type": "application/json"}
    
    all_embedded_chunks = []
    BATCH_SIZE = 100 

    progress_bar = st.progress(0, text="Embedding document chunks...")

    for i in range(0, len(_chunks), BATCH_SIZE):
        chunk_batch = _chunks[i:i + BATCH_SIZE]
        
        requests_list = [
            {"model": "models/embedding-001", "content": {"parts": [{"text": chunk}]}}
            for chunk in chunk_batch
        ]
        body = {"requests": requests_list}

        max_retries = 3
        delay = 2
        for attempt in range(max_retries):
            try:
                resp = requests.post(batch_embed_url, headers=headers, json=body, timeout=180)
                resp.raise_for_status()
                
                embeddings = resp.json().get("embeddings", [])
                if len(embeddings) != len(chunk_batch):
                    raise ValueError(f"API Error: Mismatch in batch. Sent {len(chunk_batch)}, got {len(embeddings)}.")

                batch_embedded = [
                    {"id": f"chunk_{i+j}", "text": chunk, "embedding": emb.get("values")}
                    for j, (chunk, emb) in enumerate(zip(chunk_batch, embeddings))
                    if emb.get("values")
                ]
                all_embedded_chunks.extend(batch_embedded)
                
                progress_val = min((i + BATCH_SIZE) / len(_chunks), 1.0)
                progress_bar.progress(progress_val, text=f"Embedded {min(i + BATCH_SIZE, len(_chunks))}/{len(_chunks)} chunks...")
                break

            except requests.exceptions.RequestException as e:
                if attempt < max_retries - 1:
                    st.warning(f"API request failed (attempt {attempt+1}/{max_retries}): {e}. Retrying in {delay}s...")
                    time.sleep(delay)
                    delay *= 2
                else:
                    raise RuntimeError(f"API request failed after {max_retries} attempts: {e}")
            except (KeyError, ValueError) as e:
                raise RuntimeError(f"Failed to parse API response: {e}\nResponse: {resp.text}")

    progress_bar.empty()
    return all_embedded_chunks


def generate_response(query, relevant_chunks, api_key):
    """Generates a natural language response based on the query and relevant chunks."""
    generation_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    context = "\n\n".join([f"Clause: {c['chunk']}\nScore: {c['score']:.4f}" for c in relevant_chunks])
    
    prompt = f"""
You are a professional assistant trained to review documents. Your task is to give a direct answer to a user's query based *only* on the relevant clauses provided from the document.

**User Query:**
"{query}"

**Relevant Clauses from Document:**
---
{context}
---

**Instructions:**
1.  Provide a direct, concise, natural language answer in one or two sentences.
2.  Start your answer directly with "Yes," "No," or state that the information isn't available.
3.  **Do not** use phrases like "According to the text," "The provided clauses state," or similar preambles. Answer the question as if you are the definitive source.
4.  If the information is not in the clauses, state that the answer cannot be found in the provided sections.

**Example Answer Formats:**
- "Yes, emergency treatment is covered for up to six weeks per trip."
- "No, the policy does not cover this specific situation."
"""
    
    headers = {"Content-Type": "application/json"}
    body = {"contents": [{"parts": [{"text": prompt}]}]}
    
    try:
        resp = requests.post(generation_url, headers=headers, json=body, timeout=120)
        resp.raise_for_status()
        response_data = resp.json()
        candidates = response_data.get("candidates", [])
        if candidates and "content" in candidates[0] and "parts" in candidates[0]["content"]:
            return candidates[0]["content"]["parts"][0].get("text", "Sorry, I couldn't generate a valid response.")
        else:
            return f"Sorry, the API returned an unexpected response format.\n{response_data}"
    except requests.exceptions.RequestException as e:
        return f"Sorry, I couldn't generate a response due to an API error: {e}"

# --- SEMANTIC SEARCH ---
def cosine_similarity(v1, v2):
    """Calculates the cosine similarity between two vectors."""
    v1, v2 = np.array(v1), np.array(v2)
    dot_product = np.dot(v1, v2)
    norm_v1 = np.linalg.norm(v1)
    norm_v2 = np.linalg.norm(v2)
    if norm_v1 == 0 or norm_v2 == 0:
        return 0.0
    return dot_product / (norm_v1 * norm_v2)

def search_relevant_chunks(query, embedded_chunks, api_key, top_k=3):
    """Finds the most relevant chunks to a query using cosine similarity."""
    if not embedded_chunks:
        return []
    query_embedding = get_embedding(query, api_key)
    scored_chunks = []
    for chunk in embedded_chunks:
        score = cosine_similarity(query_embedding, chunk["embedding"])
        scored_chunks.append({"chunk": chunk["text"], "score": score})
    
    return sorted(scored_chunks, key=lambda x: x["score"], reverse=True)[:top_k]

# --- REPORT GENERATION ---
def create_report_text(query, answer, chunks):
    """Creates a formatted string for the report."""
    report = f"Query:\n{query}\n\n"
    report += f"Answer:\n{answer}\n\n"
    report += "--- Relevant Clauses ---\n\n"
    for chunk in chunks:
        report += f"Score: {chunk['score']:.4f}\n"
        report += f"{chunk['chunk']}\n\n"
    return report

def create_docx_report(query, answer, chunks):
    """Creates a DOCX report in memory."""
    doc = docx.Document()
    doc.add_heading('Smart Document Query Report', 0)
    
    doc.add_heading('User Query', level=1)
    doc.add_paragraph(query)
    
    doc.add_heading('Generated Answer', level=1)
    doc.add_paragraph(answer)
    
    doc.add_heading('Relevant Clauses', level=1)
    for chunk in chunks:
        doc.add_paragraph(f"Score: {chunk['score']:.4f}", style='Intense Quote')
        doc.add_paragraph(chunk['chunk'])
        doc.add_paragraph() # Add a space
        
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- STREAMLIT UI ---
st.set_page_config(page_title="Smart Document Query", layout="wide")
st.title(" CodeBusters | Smart Document Query System")

# --- API KEY INPUT ---
st.sidebar.title("Configuration")
st.sidebar.markdown("Enter your Gemini API Key to get started.")
GEMINI_API_KEY = st.sidebar.text_input("Gemini API Key", type="password", help="Get your API key from Google AI Studio.")

if not GEMINI_API_KEY:
    st.info("Please enter your Gemini API Key in the sidebar to use the app.")
    st.stop()

# --- MAIN APP LOGIC ---
st.markdown("Upload a document (`PDF` or `DOCX`) and ask a question in natural language.")

# Initialize session state
if 'processed_data' not in st.session_state:
    st.session_state.processed_data = None
if 'file_name' not in st.session_state:
    st.session_state.file_name = None
if 'report_data' not in st.session_state:
    st.session_state.report_data = None

col1, col2 = st.columns(2)

with col1:
    st.subheader("1. Upload Your Document")
    uploaded_file = st.file_uploader("📎 Select a PDF or DOCX file", type=["pdf", "docx"])
    
    if uploaded_file:
        file_bytes = uploaded_file.getvalue()
        file_id = f"{uploaded_file.name}-{len(file_bytes)}"

        if st.session_state.file_name != file_id:
            st.session_state.report_data = None # Clear old report data
            with st.spinner(f"Reading and analyzing '{uploaded_file.name}'..."):
                temp_dir = tempfile.mkdtemp()
                temp_path = os.path.join(temp_dir, uploaded_file.name)
                with open(temp_path, "wb") as f:
                    f.write(file_bytes)

                try:
                    full_text = extract_text(temp_path, uploaded_file.type)
                    chunks = chunk_text(full_text)
                    if not chunks:
                         st.error("Could not extract any text or chunks from the document.")
                         st.stop()
                    
                    embedded_chunks = embed_chunks(tuple(chunks), GEMINI_API_KEY)
                    st.session_state.processed_data = embedded_chunks
                    st.session_state.file_name = file_id
                    st.success(f"✅ Document '{uploaded_file.name}' processed!")
                    st.info(f"Extracted {len(full_text.split())} words and created {len(chunks)} chunks.")

                except Exception as e:
                    st.error(f"❌ Error processing document: {str(e)}")
                    st.session_state.processed_data = None
                    st.session_state.file_name = None
                finally:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                    if os.path.exists(temp_dir):
                        os.rmdir(temp_dir)

with col2:
    st.subheader("2. Ask a Question")
    user_query = st.text_area("💬 Enter your query here:", placeholder="e.g., 'Is knee surgery covered under the policy?'", height=150)

    if st.button("🔍 Get Answer", disabled=not st.session_state.processed_data):
        if user_query.strip() and st.session_state.processed_data:
            with st.spinner("Searching and generating response..."):
                try:
                    top_chunks = search_relevant_chunks(user_query, st.session_state.processed_data, GEMINI_API_KEY, top_k=5)
                    
                    if not top_chunks:
                        st.warning("Could not find any relevant information for your query.")
                        st.session_state.report_data = None
                    else:
                        answer = generate_response(user_query, top_chunks, GEMINI_API_KEY)
                        
                        st.success("💡 Answer")
                        st.markdown(f"> {answer}")

                        # Store data for report download
                        st.session_state.report_data = {
                            "query": user_query,
                            "answer": answer,
                            "chunks": top_chunks
                        }

                        with st.expander("📚 View most relevant clauses found"):
                            for chunk in top_chunks:
                                st.markdown(f"**Score: {chunk['score']:.4f}**")
                                st.markdown(f"_{chunk['chunk']}_")
                                st.divider()

                except Exception as e:
                    st.error(f"❌ An error occurred during search: {str(e)}")
                    st.session_state.report_data = None
        elif not st.session_state.processed_data:
            st.warning("Please upload and process a document first.")
        else:
            st.warning("Please enter a query.")

    # --- DOWNLOAD SECTION ---
    if st.session_state.report_data:
        st.subheader("3. Download Report")
        report_data = st.session_state.report_data
        
        dl_col1, dl_col2 = st.columns(2)

        with dl_col1:
            st.download_button(
                label="📥 Download as TXT",
                data=create_report_text(report_data['query'], report_data['answer'], report_data['chunks']),
                file_name="report.txt",
                mime="text/plain"
            )
        with dl_col2:
            st.download_button(
                label="📥 Download as DOCX",
                data=create_docx_report(report_data['query'], report_data['answer'], report_data['chunks']),
                file_name="report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
